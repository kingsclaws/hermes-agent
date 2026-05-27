"""Lex-DOCX native tools — legal document review, editing, and generation.

Registers lex_docx operations as first-class Hermes agent tools.  The agent
can inspect, review, edit, and finalize .docx files without shelling out to
the CLI or going through an MCP middle layer.

Tool groups (matching the legal-review workflow):
  Inspect   — stats, structure, para_query, extract_table, table_inspect,
              tc_list, comment_list, footer_audit, numbering_inspect
  Edit      — insert, replace, delete (all support Track Changes)
  Review    — lint, doctor, review_stats
  Finalize  — tc_accept, tc_reject, comment_clean, header_clean, clean
  Format    — highlight, set_outline_level, bold_terms, format_brush,
              format_table, cleanup
  Generate  — create, new_table, fill_table, toc
"""

from __future__ import annotations

import json
import logging
import sys
from pathlib import Path

from tools.registry import invalidate_check_fn_cache, registry, tool_error, tool_result

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

# lex_docx is installed as an editable pip package from
# /root/.hermes/tools/lex-docx-build.  Make sure its parent is on sys.path
# so ``import lex_docx`` resolves at call time.
_LEX_DOCX_BUILD = Path("/root/.hermes/tools/lex-docx-build")
if str(_LEX_DOCX_BUILD) not in sys.path:
    sys.path.insert(0, str(_LEX_DOCX_BUILD))


def _check_lex_docx():
    """Return True when lex_docx is importable."""
    import importlib.util
    try:
        return importlib.util.find_spec("lex_docx") is not None
    except (ImportError, ValueError):
        return False


def _openxml_doc(path: str):
    """Open a .docx via the OpenXml backend (for TC ops, lint, etc.)."""
    from lex_docx.openxml_package import OpenXmlDocument
    return OpenXmlDocument(path)


def _pydocx_doc(path: str):
    """Open a .docx via python-docx (for doctor, format ops, etc.)."""
    from docx import Document
    return Document(path)


def _resolve_output(docx: str, output: str | None) -> str:
    """Return the path to save to (output overrides docx for in-place)."""
    return output or docx


def _save_openxml(doc, path: str) -> None:
    doc.save(path)


def _parse_range(raw: str | None) -> tuple[int, int] | None:
    """Parse 'lo,hi' into (lo, hi+1) half-open range for lex_docx."""
    if not raw:
        return None
    parts = [x.strip() for x in raw.split(",")]
    if len(parts) != 2:
        return None
    return (int(parts[0]), int(parts[1]) + 1)


# ---------------------------------------------------------------------------
# Tool schemas & handlers
# ---------------------------------------------------------------------------

# ── Inspect ────────────────────────────────────────────────────────────────

LEX_DOCX_STATS_SCHEMA = {
    "name": "lex_docx_stats",
    "description": (
        "Quick document summary: paragraph count, table count, font distribution "
        "top-5, Track Changes count, comment count.  Always call this FIRST when "
        "reviewing a document to understand its scale and structure."
    ),
    "parameters": {
        "type": "object",
        "properties": {
            "docx": {
                "type": "string",
                "description": "Path to the .docx file.",
            },
        },
        "required": ["docx"],
    },
}


def _handle_stats(args: dict, **kwargs) -> str:
    from lex_docx import doc_stats
    result = doc_stats.doc_stats(args["docx"])
    if "font_distribution" in result:
        result["fonts_top5"] = dict(
            list(result["font_distribution"].items())[:5]
        )
        del result["font_distribution"]
    return tool_result(result)


LEX_DOCX_EXPORT_STRUCTURE_SCHEMA = {
    "name": "lex_docx_export_structure",
    "description": (
        "Export the document structure as a tree: heading levels, paragraph "
        "previews, table locations, page/section breaks.  Use this to understand "
        "the document's outline before diving into specific sections."
    ),
    "parameters": {
        "type": "object",
        "properties": {
            "docx": {
                "type": "string",
                "description": "Path to the .docx file.",
            },
            "preview_len": {
                "type": "integer",
                "description": "Max characters per paragraph preview (default 120).",
            },
            "tree": {
                "type": "boolean",
                "description": "Output as indented tree instead of flat list (default true).",
            },
        },
        "required": ["docx"],
    },
}


def _handle_export_structure(args: dict, **kwargs) -> str:
    from lex_docx.iter_ops import export_structure
    result = export_structure(
        args["docx"],
        preview_len=args.get("preview_len", 120),
        tree=args.get("tree", True),
    )
    return tool_result(result)


LEX_DOCX_PARA_QUERY_SCHEMA = {
    "name": "lex_docx_para_query",
    "description": (
        "Search paragraphs by formatting properties: font name, font size, "
        "bold/italic, style name (e.g. 'Heading 1'), outline level, alignment, "
        "or text content regex.  Returns matching paragraph indices and previews.  "
        "Use this to locate specific sections, check formatting consistency, or "
        "find paragraphs that need attention."
    ),
    "parameters": {
        "type": "object",
        "properties": {
            "docx": {
                "type": "string",
                "description": "Path to the .docx file.",
            },
            "style": {
                "type": "string",
                "description": "Filter by paragraph style name (e.g. 'Heading 1', 'Normal').",
            },
            "font": {
                "type": "string",
                "description": "Filter by font name (east-Asian font).",
            },
            "font_size": {
                "type": "number",
                "description": "Filter by font size in points (e.g. 11.5).",
            },
            "bold": {
                "type": "boolean",
                "description": "Filter paragraphs whose first run is bold.",
            },
            "italic": {
                "type": "boolean",
                "description": "Filter paragraphs whose first run is italic.",
            },
            "outline_level": {
                "type": "integer",
                "description": "Filter by outline level (0-8, where 0=body text).",
            },
            "alignment": {
                "type": "string",
                "description": "Filter by text alignment: LEFT, CENTER, RIGHT, JUSTIFY.",
            },
            "text_regex": {
                "type": "string",
                "description": "Regex pattern to match against paragraph text.",
            },
            "para_range": {
                "type": "string",
                "description": "Limit search to paragraph range 'lo,hi' (0-indexed).",
            },
        },
        "required": ["docx"],
    },
}


def _handle_para_query(args: dict, **kwargs) -> str:
    from lex_docx.para_query import query
    filters = {}
    for key in ("style", "font", "font_size", "bold", "italic",
                "outline_level", "alignment", "text_regex"):
        if key in args and args[key] is not None:
            filters[key] = args[key]
    para_range = _parse_range(args.get("para_range"))
    results = query(args["docx"], filters=filters, para_range=para_range)
    return tool_result(results)


LEX_DOCX_EXTRACT_TABLE_SCHEMA = {
    "name": "lex_docx_extract_table",
    "description": (
        "Extract table data from a .docx as structured JSON.  Returns headers "
        "and row data.  Use this to pull out financials, term sheets, comparison "
        "tables, or any tabular content for review."
    ),
    "parameters": {
        "type": "object",
        "properties": {
            "docx": {
                "type": "string",
                "description": "Path to the .docx file.",
            },
            "table_index": {
                "type": "integer",
                "description": "Table index (0-based).  If omitted, extracts all tables.",
            },
            "near_text": {
                "type": "string",
                "description": "Find the table nearest to this text snippet.",
            },
        },
        "required": ["docx"],
    },
}


def _handle_extract_table(args: dict, **kwargs) -> str:
    from lex_docx.table_ops import extract_table
    kwargs_in = {}
    if "table_index" in args and args["table_index"] is not None:
        kwargs_in["table_index"] = args["table_index"]
    if "near_text" in args and args["near_text"] is not None:
        kwargs_in["near_text"] = args["near_text"]
    result = extract_table(args["docx"], **kwargs_in)
    return tool_result(result)


LEX_DOCX_TABLE_INSPECT_SCHEMA = {
    "name": "lex_docx_table_inspect",
    "description": (
        "Read full table format information: cell shading, borders, column widths, "
        "font details, alignment, and style detection.  Use this before applying "
        "format_table to understand what needs fixing."
    ),
    "parameters": {
        "type": "object",
        "properties": {
            "docx": {
                "type": "string",
                "description": "Path to the .docx file.",
            },
            "table_index": {
                "type": "integer",
                "description": "Table index (0-based).",
            },
            "near_text": {
                "type": "string",
                "description": "Find the table nearest to this text snippet.",
            },
        },
        "required": ["docx"],
    },
}


def _handle_table_inspect(args: dict, **kwargs) -> str:
    from lex_docx.table_ops import inspect_table
    kwargs_in = {}
    if "table_index" in args and args["table_index"] is not None:
        kwargs_in["table_index"] = args["table_index"]
    if "near_text" in args and args["near_text"] is not None:
        kwargs_in["near_text"] = args["near_text"]
    result = inspect_table(args["docx"], **kwargs_in)
    return tool_result(result)


LEX_DOCX_TC_LIST_SCHEMA = {
    "name": "lex_docx_tc_list",
    "description": (
        "List all Track Changes (insertions and deletions) in the document.  "
        "Returns change ID, type (ins/del), author, and text preview for each.  "
        "Use this to review pending changes before accepting or rejecting."
    ),
    "parameters": {
        "type": "object",
        "properties": {
            "docx": {
                "type": "string",
                "description": "Path to the .docx file.",
            },
            "author": {
                "type": "string",
                "description": "Filter by revision author name.",
            },
            "tc_type": {
                "type": "string",
                "description": "Filter by type: 'ins' or 'del'.",
            },
            "para_range": {
                "type": "string",
                "description": "Limit to paragraph range 'lo,hi' (0-indexed).",
            },
        },
        "required": ["docx"],
    },
}


def _handle_tc_list(args: dict, **kwargs) -> str:
    from lex_docx import tc_ops
    doc = _openxml_doc(args["docx"])
    items = tc_ops.list_tc(
        doc,
        author_filter=args.get("author"),
        para_range=_parse_range(args.get("para_range")),
        type_filter=args.get("tc_type"),
    )
    return tool_result({"total": len(items), "items": items})


LEX_DOCX_COMMENT_LIST_SCHEMA = {
    "name": "lex_docx_comment_list",
    "description": (
        "List all comments (annotations) in the document.  Returns comment ID, "
        "author, text, and associated paragraph.  Use this alongside tc_list "
        "for a complete picture of pending review feedback."
    ),
    "parameters": {
        "type": "object",
        "properties": {
            "docx": {
                "type": "string",
                "description": "Path to the .docx file.",
            },
            "author": {
                "type": "string",
                "description": "Filter by comment author name.",
            },
            "para_range": {
                "type": "string",
                "description": "Limit to paragraph range 'lo,hi' (0-indexed).",
            },
        },
        "required": ["docx"],
    },
}


def _handle_comment_list(args: dict, **kwargs) -> str:
    from lex_docx import tc_ops
    doc = _openxml_doc(args["docx"])
    items = tc_ops.list_comments(
        doc,
        author_filter=args.get("author"),
        para_range=_parse_range(args.get("para_range")),
    )
    return tool_result({"total": len(items), "items": items})


LEX_DOCX_FOOTER_AUDIT_SCHEMA = {
    "name": "lex_docx_footer_audit",
    "description": (
        "Audit all footer content in the document.  Returns footer text per "
        "section.  Use this to check for stale entity names, outdated dates, "
        "or template residue in footers."
    ),
    "parameters": {
        "type": "object",
        "properties": {
            "docx": {
                "type": "string",
                "description": "Path to the .docx file.",
            },
        },
        "required": ["docx"],
    },
}


def _handle_footer_audit(args: dict, **kwargs) -> str:
    from lex_docx.footer_ops import audit_footers
    doc = _openxml_doc(args["docx"])
    items = audit_footers(doc)
    return tool_result({"total": len(items), "items": items})


LEX_DOCX_NUMBERING_INSPECT_SCHEMA = {
    "name": "lex_docx_numbering_inspect",
    "description": (
        "Inspect paragraph numbering state: which paragraphs have numbering, "
        "their numId/ilvl, and whether the numbering is own or inherited from "
        "a style.  Use this to diagnose numbering gaps or inconsistencies."
    ),
    "parameters": {
        "type": "object",
        "properties": {
            "docx": {
                "type": "string",
                "description": "Path to the .docx file.",
            },
            "para_range": {
                "type": "string",
                "description": "Limit to paragraph range 'lo,hi' (0-indexed).",
            },
        },
        "required": ["docx"],
    },
}


def _handle_numbering_inspect(args: dict, **kwargs) -> str:
    from lex_docx.numbering_ops import inspect_numbering
    doc = _pydocx_doc(args["docx"])
    para_range = _parse_range(args.get("para_range"))
    items = inspect_numbering(doc, para_range=para_range)
    return tool_result({"total": len(items), "items": items})


# ── Edit ───────────────────────────────────────────────────────────────────

LEX_DOCX_INSERT_SCHEMA = {
    "name": "lex_docx_insert",
    "description": (
        "Insert text at the end of a paragraph.  When tc=True, the insertion "
        "is marked as a Track Change so it can be reviewed before finalizing.  "
        "ALWAYS use tc=True for legal document review — never silently modify text."
    ),
    "parameters": {
        "type": "object",
        "properties": {
            "docx": {
                "type": "string",
                "description": "Path to the .docx file.",
            },
            "para": {
                "type": "integer",
                "description": "Paragraph index (0-based).",
            },
            "text": {
                "type": "string",
                "description": "Text to insert.",
            },
            "tc": {
                "type": "boolean",
                "description": "Enable Track Changes mode (default true for legal review).",
            },
            "author": {
                "type": "string",
                "description": "Revision author name (default 'agent').",
            },
            "output": {
                "type": "string",
                "description": "Output path (default: overwrite input).",
            },
        },
        "required": ["docx", "para", "text"],
    },
}


def _handle_insert(args: dict, **kwargs) -> str:
    from lex_docx.edit_ops import insert_text
    result = insert_text(
        args["docx"], args["para"], args["text"],
        tc=args.get("tc", True),
        author=args.get("author", "agent"),
        output=args.get("output"),
    )
    return tool_result(result.__dict__)


LEX_DOCX_REPLACE_SCHEMA = {
    "name": "lex_docx_replace",
    "description": (
        "Replace text within a paragraph.  When tc=True, the old text is marked "
        "as deleted and the new text as inserted via Track Changes.  "
        "ALWAYS use tc=True for legal document review."
    ),
    "parameters": {
        "type": "object",
        "properties": {
            "docx": {
                "type": "string",
                "description": "Path to the .docx file.",
            },
            "para": {
                "type": "integer",
                "description": "Paragraph index (0-based).",
            },
            "old": {
                "type": "string",
                "description": "Exact text to replace.",
            },
            "new": {
                "type": "string",
                "description": "Replacement text.",
            },
            "tc": {
                "type": "boolean",
                "description": "Enable Track Changes mode (default true).",
            },
            "author": {
                "type": "string",
                "description": "Revision author name (default 'agent').",
            },
            "output": {
                "type": "string",
                "description": "Output path (default: overwrite input).",
            },
        },
        "required": ["docx", "para", "old", "new"],
    },
}


def _handle_replace(args: dict, **kwargs) -> str:
    from lex_docx.edit_ops import replace_text
    result = replace_text(
        args["docx"], args["para"], args["old"], args["new"],
        tc=args.get("tc", True),
        author=args.get("author", "agent"),
        output=args.get("output"),
    )
    return tool_result(result.__dict__)


LEX_DOCX_DELETE_SCHEMA = {
    "name": "lex_docx_delete",
    "description": (
        "Delete text within a paragraph.  When tc=True, the deletion is marked "
        "as a Track Change.  If text is omitted, the entire paragraph is marked "
        "for deletion.  ALWAYS use tc=True for legal document review."
    ),
    "parameters": {
        "type": "object",
        "properties": {
            "docx": {
                "type": "string",
                "description": "Path to the .docx file.",
            },
            "para": {
                "type": "integer",
                "description": "Paragraph index (0-based).",
            },
            "text": {
                "type": "string",
                "description": "Exact text to delete.  Omit to delete entire paragraph.",
            },
            "tc": {
                "type": "boolean",
                "description": "Enable Track Changes mode (default true).",
            },
            "author": {
                "type": "string",
                "description": "Revision author name (default 'agent').",
            },
            "output": {
                "type": "string",
                "description": "Output path (default: overwrite input).",
            },
        },
        "required": ["docx", "para"],
    },
}


def _handle_delete(args: dict, **kwargs) -> str:
    from lex_docx.edit_ops import delete_text, delete_paragraph_tc
    if args.get("text"):
        result = delete_text(
            args["docx"], args["para"], text=args["text"],
            tc=args.get("tc", True),
            author=args.get("author", "agent"),
            output=args.get("output"),
        )
    else:
        result = delete_paragraph_tc(
            args["docx"], args["para"],
            author=args.get("author", "agent"),
            output=args.get("output"),
        )
    return tool_result(result.__dict__)


# ── Review ─────────────────────────────────────────────────────────────────

LEX_DOCX_LINT_SCHEMA = {
    "name": "lex_docx_lint",
    "description": (
        "Check document formatting against rules: JT Note format, forbidden text "
        "patterns, entity name consistency, TC author uniformity, indent "
        "consistency, defined terms bold, table header format, table borders, "
        "spelling.  Returns per-rule pass/fail with locations.  Use this early "
        "in review to catch formatting issues before detailed content review."
    ),
    "parameters": {
        "type": "object",
        "properties": {
            "docx": {
                "type": "string",
                "description": "Path to the .docx file.",
            },
            "rules": {
                "type": "string",
                "description": "Comma-separated rule names to run (default: all).",
            },
            "para_range": {
                "type": "string",
                "description": "Limit check to paragraph range 'lo,hi' (0-indexed).",
            },
        },
        "required": ["docx"],
    },
}


def _handle_lint(args: dict, **kwargs) -> str:
    from lex_docx import lint, DocConfig
    rule_list = args["rules"].split(",") if args.get("rules") else None
    results = lint.check(
        args["docx"],
        config=DocConfig(),
        rules=rule_list,
    )
    return tool_result({
        "total": len(results),
        "passed": sum(1 for r in results if r.passed),
        "failed": sum(1 for r in results if not r.passed),
        "results": [
            {
                "rule": r.rule,
                "severity": r.severity,
                "passed": r.passed,
                "detail": r.detail,
                "locations": r.locations[:20],
            }
            for r in results
        ],
    })


LEX_DOCX_DOCTOR_SCHEMA = {
    "name": "lex_docx_doctor",
    "description": (
        "Format diagnostics for the document.  When action='check' (default), "
        "returns a list of format issues: missing/mismatched fonts (D01/D02), "
        "double numbering (D03), outline leaks (D04), numbering gaps (D05), "
        "invalid style refs (D06), TOC issues (D07), heading font inconsistency "
        "(D08), footer stale entities (D09).  When action='fix', auto-repairs "
        "D01/D02/D04/D05/D07/D08 (D03/D06/D09 require human judgment).  "
        "Use 'check' first, review the issues, then 'fix'."
    ),
    "parameters": {
        "type": "object",
        "properties": {
            "docx": {
                "type": "string",
                "description": "Path to the .docx file.",
            },
            "action": {
                "type": "string",
                "description": "'check' (default) or 'fix'.",
            },
            "para_range": {
                "type": "string",
                "description": "Limit to paragraph range 'lo,hi' (0-indexed).",
            },
            "output": {
                "type": "string",
                "description": "Output path (for 'fix' action; default: overwrite input).",
            },
            "dry_run": {
                "type": "boolean",
                "description": "Preview fixes without saving (for 'fix' action).",
            },
        },
        "required": ["docx"],
    },
}


def _handle_doctor(args: dict, **kwargs) -> str:
    from lex_docx import doctor as dr
    action = args.get("action", "check")
    doc = _pydocx_doc(args["docx"])
    para_range = _parse_range(args.get("para_range"))

    if action == "fix":
        check_result = dr.check(doc, para_range=para_range)
        fix_result = dr.fix(
            doc, check_result.issues,
            dry_run=args.get("dry_run", False),
        )
        if not args.get("dry_run"):
            _save_openxml(doc, _resolve_output(args["docx"], args.get("output")))
        return tool_result({
            "action": "fix",
            "dry_run": args.get("dry_run", False),
            **fix_result,
        })

    result = dr.check(doc, para_range=para_range)
    return tool_result({
        "action": "check",
        "total_issues": len(result.issues),
        "issues": [
            {
                "rule": i.rule,
                "para": i.para,
                "detail": i.detail,
                "auto_fixable": i.rule in ("D01", "D02", "D04", "D05", "D07", "D08"),
            }
            for i in result.issues
        ],
    })


LEX_DOCX_REVIEW_STATS_SCHEMA = {
    "name": "lex_docx_review_stats",
    "description": (
        "Pre-cleanup review summary: counts Track Changes (by author), comments, "
        "footer issues, and D09 footer stale-entity findings.  Use this before "
        "'lex_docx_clean' to understand what will be affected."
    ),
    "parameters": {
        "type": "object",
        "properties": {
            "docx": {
                "type": "string",
                "description": "Path to the .docx file.",
            },
            "author": {
                "type": "string",
                "description": "Filter TC/comments by author.",
            },
            "para_range": {
                "type": "string",
                "description": "Limit to paragraph range 'lo,hi' (0-indexed).",
            },
        },
        "required": ["docx"],
    },
}


def _handle_review_stats(args: dict, **kwargs) -> str:
    from lex_docx import tc_ops, doctor as dr, footer_ops
    ox_doc = _openxml_doc(args["docx"])
    doc = _pydocx_doc(args["docx"])
    para_range = _parse_range(args.get("para_range"))

    tc_items = tc_ops.list_tc(
        ox_doc,
        author_filter=args.get("author"),
        para_range=para_range,
    )
    comments = tc_ops.list_comments(
        ox_doc,
        author_filter=args.get("author"),
        para_range=para_range,
    )
    footers = footer_ops.audit_footers(ox_doc)

    check_result = dr.check(doc, para_range=para_range)
    d09_issues = [i for i in check_result.issues if i.rule == "D09"]

    # Count TCs by author
    tc_by_author = {}
    for item in tc_items:
        a = item.get("author", "unknown")
        tc_by_author[a] = tc_by_author.get(a, 0) + 1

    return tool_result({
        "tc_total": len(tc_items),
        "tc_by_author": tc_by_author,
        "comments_total": len(comments),
        "footer_parts": len(footers),
        "footer_issues": [
            {"section": f.get("section"), "text_preview": f.get("text", "")[:200]}
            for f in footers if f.get("text")
        ],
        "d09_stale_entities": len(d09_issues),
        "d09_details": [i.detail for i in d09_issues[:10]],
    })


# ── Finalize ───────────────────────────────────────────────────────────────

LEX_DOCX_TC_ACCEPT_SCHEMA = {
    "name": "lex_docx_tc_accept",
    "description": (
        "Accept Track Changes in the document.  Optionally filter by author, "
        "change type (ins/del), or paragraph range.  Use --dry-run first to "
        "preview what will be accepted."
    ),
    "parameters": {
        "type": "object",
        "properties": {
            "docx": {
                "type": "string",
                "description": "Path to the .docx file.",
            },
            "author": {
                "type": "string",
                "description": "Accept only changes by this author.",
            },
            "tc_type": {
                "type": "string",
                "description": "Accept only 'ins' or 'del' changes.",
            },
            "para_range": {
                "type": "string",
                "description": "Limit to paragraph range 'lo,hi' (0-indexed).",
            },
            "output": {
                "type": "string",
                "description": "Output path (default: overwrite input).",
            },
            "dry_run": {
                "type": "boolean",
                "description": "Preview without modifying.",
            },
        },
        "required": ["docx"],
    },
}


def _handle_tc_accept(args: dict, **kwargs) -> str:
    from lex_docx import tc_ops
    doc = _openxml_doc(args["docx"])
    para_range = _parse_range(args.get("para_range"))

    if args.get("dry_run"):
        items = tc_ops.list_tc(
            doc,
            author_filter=args.get("author"),
            para_range=para_range,
            type_filter=args.get("tc_type"),
        )
        return tool_result({"dry_run": True, "would_accept": len(items), "items": items[:50]})

    stats = tc_ops.accept_all(
        doc,
        author_filter=args.get("author"),
        para_range=para_range,
        type_filter=args.get("tc_type"),
    )
    _save_openxml(doc, _resolve_output(args["docx"], args.get("output")))
    return tool_result({"ok": True, **stats})


LEX_DOCX_TC_REJECT_SCHEMA = {
    "name": "lex_docx_tc_reject",
    "description": (
        "Reject (revert) Track Changes in the document.  Optionally filter by "
        "author, change type, or paragraph range.  Use --dry-run first."
    ),
    "parameters": {
        "type": "object",
        "properties": {
            "docx": {
                "type": "string",
                "description": "Path to the .docx file.",
            },
            "author": {
                "type": "string",
                "description": "Reject only changes by this author.",
            },
            "tc_type": {
                "type": "string",
                "description": "Reject only 'ins' or 'del' changes.",
            },
            "para_range": {
                "type": "string",
                "description": "Limit to paragraph range 'lo,hi' (0-indexed).",
            },
            "output": {
                "type": "string",
                "description": "Output path (default: overwrite input).",
            },
            "dry_run": {
                "type": "boolean",
                "description": "Preview without modifying.",
            },
        },
        "required": ["docx"],
    },
}


def _handle_tc_reject(args: dict, **kwargs) -> str:
    from lex_docx import tc_ops
    doc = _openxml_doc(args["docx"])
    para_range = _parse_range(args.get("para_range"))

    if args.get("dry_run"):
        items = tc_ops.list_tc(
            doc,
            author_filter=args.get("author"),
            para_range=para_range,
            type_filter=args.get("tc_type"),
        )
        return tool_result({"dry_run": True, "would_reject": len(items), "items": items[:50]})

    stats = tc_ops.reject_all(
        doc,
        author_filter=args.get("author"),
        para_range=para_range,
        type_filter=args.get("tc_type"),
    )
    _save_openxml(doc, _resolve_output(args["docx"], args.get("output")))
    return tool_result({"ok": True, **stats})


LEX_DOCX_COMMENT_CLEAN_SCHEMA = {
    "name": "lex_docx_comment_clean",
    "description": (
        "Remove comments (annotations) from the document.  Optionally filter "
        "by author or paragraph range.  Use --dry-run first to preview."
    ),
    "parameters": {
        "type": "object",
        "properties": {
            "docx": {
                "type": "string",
                "description": "Path to the .docx file.",
            },
            "author": {
                "type": "string",
                "description": "Remove only comments by this author.",
            },
            "para_range": {
                "type": "string",
                "description": "Limit to paragraph range 'lo,hi' (0-indexed).",
            },
            "output": {
                "type": "string",
                "description": "Output path (default: overwrite input).",
            },
            "dry_run": {
                "type": "boolean",
                "description": "Preview without modifying.",
            },
        },
        "required": ["docx"],
    },
}


def _handle_comment_clean(args: dict, **kwargs) -> str:
    from lex_docx import tc_ops
    doc = _openxml_doc(args["docx"])
    para_range = _parse_range(args.get("para_range"))

    if args.get("dry_run"):
        items = tc_ops.list_comments(
            doc,
            author_filter=args.get("author"),
            para_range=para_range,
        )
        return tool_result({"dry_run": True, "total": len(items), "items": items[:50]})

    stats = tc_ops.clean_comments_filtered(
        doc,
        author_filter=args.get("author"),
        para_range=para_range,
    )
    _save_openxml(doc, _resolve_output(args["docx"], args.get("output")))
    return tool_result({"ok": True, **stats})


LEX_DOCX_HEADER_CLEAN_SCHEMA = {
    "name": "lex_docx_header_clean",
    "description": (
        "Clear all header content in the document.  Optionally also remove "
        "header references from sections.  Use --dry-run first to preview."
    ),
    "parameters": {
        "type": "object",
        "properties": {
            "docx": {
                "type": "string",
                "description": "Path to the .docx file.",
            },
            "remove_refs": {
                "type": "boolean",
                "description": "Also remove headerReference from section properties.",
            },
            "output": {
                "type": "string",
                "description": "Output path (default: overwrite input).",
            },
            "dry_run": {
                "type": "boolean",
                "description": "Preview without modifying.",
            },
        },
        "required": ["docx"],
    },
}


def _handle_header_clean(args: dict, **kwargs) -> str:
    from lex_docx import tc_ops
    doc = _openxml_doc(args["docx"])

    if args.get("dry_run"):
        header_count = sum(
            1 for rel in doc.part.rels.values()
            if tc_ops._HEADERS_REL in rel.reltype
        )
        return tool_result({"dry_run": True, "header_parts_found": header_count})

    stats = tc_ops.clean_headers(
        doc,
        clear_text=True,
        remove_refs=args.get("remove_refs", False),
    )
    _save_openxml(doc, _resolve_output(args["docx"], args.get("output")))
    return tool_result({"ok": True, **stats})


LEX_DOCX_CLEAN_SCHEMA = {
    "name": "lex_docx_clean",
    "description": (
        "One-click execution-version cleanup: (1) accept or reject all Track "
        "Changes, (2) remove all comments, (3) clear all header content.  "
        "This is the final step before issuing an execution version.  "
        "ALWAYS run lex_docx_review_stats first to preview what will be affected."
    ),
    "parameters": {
        "type": "object",
        "properties": {
            "docx": {
                "type": "string",
                "description": "Path to the .docx file.",
            },
            "tc_mode": {
                "type": "string",
                "description": "'accept' (default) or 'reject' all Track Changes.",
            },
            "author": {
                "type": "string",
                "description": "Filter TC/comments by author.",
            },
            "keep_comments": {
                "type": "boolean",
                "description": "Skip comment removal.",
            },
            "keep_headers": {
                "type": "boolean",
                "description": "Skip header clearing.",
            },
            "output": {
                "type": "string",
                "description": "Output path (default: overwrite input).",
            },
            "dry_run": {
                "type": "boolean",
                "description": "Preview plan without modifying.",
            },
        },
        "required": ["docx"],
    },
}


def _handle_clean(args: dict, **kwargs) -> str:
    from lex_docx import tc_ops
    doc = _openxml_doc(args["docx"])
    tc_mode = args.get("tc_mode", "accept")
    author_filter = args.get("author")
    do_comments = not args.get("keep_comments", False)
    do_headers = not args.get("keep_headers", False)

    if args.get("dry_run"):
        tc_items = tc_ops.list_tc(doc, author_filter=author_filter)
        comments = tc_ops.list_comments(doc, author_filter=author_filter)
        header_count = sum(
            1 for rel in doc.part.rels.values()
            if tc_ops._HEADERS_REL in rel.reltype
        )
        return tool_result({
            "dry_run": True,
            "tc_mode": tc_mode,
            "author_filter": author_filter,
            "tc_changes_found": len(tc_items),
            "tc_changes_preview": tc_items[:20],
            "comments_found": len(comments),
            "comments_preview": comments[:20],
            "header_parts_found": header_count,
            "do_comments": do_comments,
            "do_headers": do_headers,
        })

    result: dict = {"ok": True, "tc_mode": tc_mode}
    if tc_mode == "accept":
        result["tc"] = tc_ops.accept_all(doc, author_filter=author_filter)
    else:
        result["tc"] = tc_ops.reject_all(doc, author_filter=author_filter)
    if do_comments:
        result["comments"] = tc_ops.clean_comments(doc)
    if do_headers:
        result["headers"] = tc_ops.clean_headers(doc)

    _save_openxml(doc, _resolve_output(args["docx"], args.get("output")))
    return tool_result(result)


# ── Format ─────────────────────────────────────────────────────────────────

LEX_DOCX_HIGHLIGHT_SCHEMA = {
    "name": "lex_docx_highlight",
    "description": (
        "Highlight paragraph ranges with a background color (default: yellow).  "
        "Use this to mark sections that need client attention, internal follow-up, "
        "or cross-reference checks during review."
    ),
    "parameters": {
        "type": "object",
        "properties": {
            "docx": {
                "type": "string",
                "description": "Path to the .docx file.",
            },
            "para": {
                "type": "integer",
                "description": "Single paragraph index to highlight.",
            },
            "para_range": {
                "type": "string",
                "description": "Range of paragraphs to highlight 'lo,hi' (0-indexed).",
            },
            "color": {
                "type": "string",
                "description": "Highlight color: yellow (default), cyan, magenta, green, red.",
            },
            "output": {
                "type": "string",
                "description": "Output path (default: overwrite input).",
            },
        },
        "required": ["docx"],
    },
}


def _handle_highlight(args: dict, **kwargs) -> str:
    from lxml import etree
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = _openxml_doc(args["docx"])
    color = args.get("color", "yellow")

    if args.get("para_range"):
        a, b = [int(x.strip()) for x in args["para_range"].split(",")]
        indices = list(range(a, b + 1))
    elif args.get("para") is not None:
        indices = [args["para"]]
    else:
        return tool_error("para or para_range is required")

    marked = []
    for idx in indices:
        if idx >= len(doc.paragraphs):
            continue
        para_el = doc.paragraphs[idx]._element
        runs = para_el.findall(qn("w:r"))
        if not runs:
            r = OxmlElement("w:r")
            rPr = OxmlElement("w:rPr")
            hl = OxmlElement("w:highlight")
            hl.set(qn("w:val"), color)
            rPr.append(hl)
            r.insert(0, rPr)
            para_el.append(r)
        else:
            for run_el in runs:
                rPr = run_el.find(qn("w:rPr"))
                if rPr is None:
                    rPr = OxmlElement("w:rPr")
                    run_el.insert(0, rPr)
                existing = rPr.find(qn("w:highlight"))
                if existing is not None:
                    rPr.remove(existing)
                hl = OxmlElement("w:highlight")
                hl.set(qn("w:val"), color)
                rPr.append(hl)
        marked.append(idx)

    _save_openxml(doc, _resolve_output(args["docx"], args.get("output")))
    return tool_result({"ok": True, "highlighted": marked, "color": color})


LEX_DOCX_SET_OUTLINE_LEVEL_SCHEMA = {
    "name": "lex_docx_set_outline_level",
    "description": (
        "Set the outline level (0-8) for one or more paragraphs.  Level 0 is "
        "body text; levels 1-8 correspond to heading levels.  Can target by "
        "paragraph index, range, or style name.  Use this to fix heading "
        "hierarchy so TOC generation works correctly."
    ),
    "parameters": {
        "type": "object",
        "properties": {
            "docx": {
                "type": "string",
                "description": "Path to the .docx file.",
            },
            "para": {
                "type": "integer",
                "description": "Single paragraph index.",
            },
            "para_range": {
                "type": "string",
                "description": "Range of paragraphs 'lo,hi' (0-indexed).",
            },
            "style": {
                "type": "string",
                "description": "Target all paragraphs with this style name.",
            },
            "level": {
                "type": "integer",
                "description": "Outline level to set (0=body, 1=Heading1, 2=Heading2, etc.).",
            },
            "output": {
                "type": "string",
                "description": "Output path (default: overwrite input).",
            },
        },
        "required": ["docx", "level"],
    },
}


def _handle_set_outline_level(args: dict, **kwargs) -> str:
    from lex_docx.format_brush import set_outline_level
    kwargs_in = {"level": args["level"]}
    if args.get("para") is not None:
        kwargs_in["para"] = args["para"]
    if args.get("para_range"):
        a, b = [int(x.strip()) for x in args["para_range"].split(",")]
        kwargs_in["para_range"] = (a, b)
    if args.get("style"):
        kwargs_in["style"] = args["style"]

    doc = _pydocx_doc(args["docx"])
    result = set_outline_level(doc, **kwargs_in)
    _save_openxml(doc, _resolve_output(args["docx"], args.get("output")))
    return tool_result({"ok": True, **result})


LEX_DOCX_BOLD_TERMS_SCHEMA = {
    "name": "lex_docx_bold_terms",
    "description": (
        "Auto-detect and bold defined terms in a paragraph.  Use --scan to "
        "find all candidate defined-term paragraphs first, then apply to "
        "specific paragraphs.  Common in legal documents where defined terms "
        "should be bold on first use."
    ),
    "parameters": {
        "type": "object",
        "properties": {
            "docx": {
                "type": "string",
                "description": "Path to the .docx file.",
            },
            "para": {
                "type": "integer",
                "description": "Paragraph index to process (0-based).",
            },
            "scan": {
                "type": "boolean",
                "description": "Scan the document for defined-term paragraphs instead of bolding.",
            },
            "para_range": {
                "type": "string",
                "description": "Limit scan to paragraph range 'lo,hi' (0-indexed).",
            },
            "output": {
                "type": "string",
                "description": "Output path (default: overwrite input).",
            },
        },
        "required": ["docx"],
    },
}


def _handle_bold_terms(args: dict, **kwargs) -> str:
    from lex_docx import defined_terms
    doc = _pydocx_doc(args["docx"])

    if args.get("scan"):
        para_range = _parse_range(args.get("para_range"))
        results = defined_terms.scan_terms(doc, para_range=para_range)
        return tool_result(results)

    if args.get("para") is None:
        return tool_error("para is required (or use scan=True to find candidates)")

    terms = defined_terms.auto_bold(doc, paragraph_index=args["para"])
    _save_openxml(doc, _resolve_output(args["docx"], args.get("output")))
    return tool_result({"bolded": terms})


LEX_DOCX_FORMAT_BRUSH_SCHEMA = {
    "name": "lex_docx_format_brush",
    "description": (
        "Copy paragraph formatting (font, size, bold/italic, alignment, "
        "indentation, spacing) from a reference paragraph to one or more "
        "target paragraphs.  Use this to fix inconsistent formatting across "
        "sections."
    ),
    "parameters": {
        "type": "object",
        "properties": {
            "docx": {
                "type": "string",
                "description": "Path to the .docx file.",
            },
            "ref": {
                "type": "integer",
                "description": "Reference paragraph index (source of formatting).",
            },
            "target": {
                "type": "string",
                "description": "Comma-separated target paragraph indices.",
            },
            "target_range": {
                "type": "string",
                "description": "Target paragraph range 'lo,hi' (alternative to --target).",
            },
            "copy": {
                "type": "string",
                "description": "Comma-separated properties to copy: indent,spacing,font,size,bold,italic,alignment,style.",
            },
            "output": {
                "type": "string",
                "description": "Output path (default: overwrite input).",
            },
        },
        "required": ["docx", "ref"],
    },
}


def _handle_format_brush(args: dict, **kwargs) -> str:
    from lex_docx.backends.format_backend import apply_format_brush

    if args.get("target"):
        indices = [int(x) for x in args["target"].split(",")]
    elif args.get("target_range"):
        a, b = [int(x.strip()) for x in args["target_range"].split(",")]
        indices = list(range(a, b + 1))
    else:
        return tool_error("target or target_range is required")

    copy = args["copy"].split(",") if args.get("copy") else None
    doc = _openxml_doc(args["docx"])

    modified = apply_format_brush(
        backend="ooxml",
        doc=doc,
        target_indices=indices,
        reference_index=args["ref"],
        copy=copy,
        safe=True,
    )

    if modified.get("failed_at") is not None:
        return tool_result({
            "ok": False,
            "modified": modified["modified"],
            "details": modified.get("details", []),
            "failed_at": modified["failed_at"],
            "error": modified.get("error"),
        })

    _save_openxml(doc, _resolve_output(args["docx"], args.get("output")))
    return tool_result({
        "ok": True,
        "modified": modified["modified"],
        "details": modified.get("details", []),
    })


LEX_DOCX_CLEANUP_SCHEMA = {
    "name": "lex_docx_cleanup",
    "description": (
        "Remove empty paragraphs and orphan numbering from the document.  "
        "Empty paragraphs are marked as Track Changes deletions (default) "
        "so the cleanup is reviewable.  Use this to tidy up documents "
        "before final review."
    ),
    "parameters": {
        "type": "object",
        "properties": {
            "docx": {
                "type": "string",
                "description": "Path to the .docx file.",
            },
            "para_range": {
                "type": "string",
                "description": "Limit to paragraph range 'lo,hi' (0-indexed).",
            },
            "mode": {
                "type": "string",
                "description": "'report' to only list issues, 'fix' (default) to apply TC deletions.",
            },
            "output": {
                "type": "string",
                "description": "Output path (default: overwrite input).",
            },
        },
        "required": ["docx"],
    },
}


def _handle_cleanup(args: dict, **kwargs) -> str:
    from lex_docx import cleanup
    doc = _pydocx_doc(args["docx"])
    para_range = _parse_range(args.get("para_range"))
    mode = args.get("mode", "fix")

    result = cleanup.cleanup_all(
        doc,
        as_tc_del=True,
        para_range=para_range,
    )

    if mode != "report":
        _save_openxml(doc, _resolve_output(args["docx"], args.get("output")))

    return tool_result(result)


LEX_DOCX_FORMAT_TABLE_SCHEMA = {
    "name": "lex_docx_format_table",
    "description": (
        "Apply unified formatting to a table: header row shading, borders, "
        "column widths, cell alignment, font size/bold.  Use this after "
        "table_inspect to fix formatting issues in a single pass."
    ),
    "parameters": {
        "type": "object",
        "properties": {
            "docx": {
                "type": "string",
                "description": "Path to the .docx file.",
            },
            "table_index": {
                "type": "integer",
                "description": "Table index (0-based).",
            },
            "near_text": {
                "type": "string",
                "description": "Find the table nearest to this text snippet.",
            },
            "header_bg": {
                "type": "string",
                "description": "Header row background color in hex (e.g. '1F4E79').",
            },
            "border_style": {
                "type": "string",
                "description": "Border style: single, double, none, dashed.",
            },
            "col_widths": {
                "type": "string",
                "description": "Comma-separated column widths in twips (e.g. '2000,6000,2000').",
            },
            "output": {
                "type": "string",
                "description": "Output path (default: overwrite input).",
            },
        },
        "required": ["docx"],
    },
}


def _handle_format_table(args: dict, **kwargs) -> str:
    from lex_docx.table_ops import format_table
    kwargs_in = {}
    if "table_index" in args and args["table_index"] is not None:
        kwargs_in["table_index"] = args["table_index"]
    if "near_text" in args and args["near_text"] is not None:
        kwargs_in["near_text"] = args["near_text"]
    if "header_bg" in args and args["header_bg"] is not None:
        kwargs_in["header_bg"] = args["header_bg"]
    if "border_style" in args and args["border_style"] is not None:
        kwargs_in["border_style"] = args["border_style"]
    if "col_widths" in args and args["col_widths"] is not None:
        kwargs_in["col_widths"] = [int(w) for w in args["col_widths"].split(",")]

    result = format_table(args["docx"], **kwargs_in)
    return tool_result(result)


# ── Generate ───────────────────────────────────────────────────────────────

LEX_DOCX_CREATE_SCHEMA = {
    "name": "lex_docx_create",
    "description": (
        "Create a new OPC-compliant .docx skeleton from scratch.  Sets up "
        "standard styles, fonts (Song/TNR), and metadata.  Use this to "
        "bootstrap a new legal document before filling in content."
    ),
    "parameters": {
        "type": "object",
        "properties": {
            "output": {
                "type": "string",
                "description": "Output .docx file path.",
            },
            "title": {
                "type": "string",
                "description": "Document title.",
            },
            "meta": {
                "type": "string",
                "description": "Metadata lines joined by newline (case number, date, etc.).",
            },
            "font_size": {
                "type": "number",
                "description": "Body font size in pt (default 11.5).",
            },
        },
        "required": ["output"],
    },
}


def _handle_create(args: dict, **kwargs) -> str:
    from lex_docx.doc_create import create_document
    result = create_document(
        output=args["output"],
        title=args.get("title", ""),
        meta=args.get("meta", ""),
        font_size=args.get("font_size", 11.5),
    )
    return tool_result(result)


LEX_DOCX_NEW_TABLE_SCHEMA = {
    "name": "lex_docx_new_table",
    "description": (
        "Insert a new table into the document.  Supports grid tables, key-value "
        "tables, merged-cell tables, nested tables, and diagonal-header tables.  "
        "Use after 'create' or to add comparison tables to an existing document."
    ),
    "parameters": {
        "type": "object",
        "properties": {
            "docx": {
                "type": "string",
                "description": "Path to the .docx file.",
            },
            "table_type": {
                "type": "string",
                "description": "Table type: grid, kv, merged, nested, diagonal (default: grid).",
            },
            "cols": {
                "type": "integer",
                "description": "Number of columns (for grid type).",
            },
            "rows": {
                "type": "integer",
                "description": "Number of data rows (excluding header).",
            },
            "headers": {
                "type": "string",
                "description": "JSON array of header texts, e.g. '[\"Column A\",\"Column B\"]'.",
            },
            "data": {
                "type": "string",
                "description": "JSON array of row arrays, e.g. '[[\"a1\",\"b1\"],[\"a2\",\"b2\"]]'.",
            },
            "after_para": {
                "type": "integer",
                "description": "Insert the table after this paragraph index.",
            },
            "output": {
                "type": "string",
                "description": "Output path (default: overwrite input).",
            },
        },
        "required": ["docx"],
    },
}


def _handle_new_table(args: dict, **kwargs) -> str:
    from lex_docx.new_table_ops import insert_table
    kwargs_in = {"type": args.get("table_type", "grid")}
    if "cols" in args and args["cols"] is not None:
        kwargs_in["cols"] = args["cols"]
    if "rows" in args and args["rows"] is not None:
        kwargs_in["rows"] = args["rows"]
    if "headers" in args and args["headers"] is not None:
        kwargs_in["headers"] = json.loads(args["headers"])
    if "data" in args and args["data"] is not None:
        kwargs_in["data"] = json.loads(args["data"])
    if "after_para" in args and args["after_para"] is not None:
        kwargs_in["at"] = args["after_para"]

    result = insert_table(args["docx"], **kwargs_in)
    return tool_result(result)


LEX_DOCX_FILL_TABLE_SCHEMA = {
    "name": "lex_docx_fill_table",
    "description": (
        "Fill a table by column mapping.  Provide a JSON mapping of column "
        "names to values, and the tool locates the correct row and writes "
        "values into the corresponding cells.  Supports fuzzy column name "
        "matching.  Use this to populate template tables with review data."
    ),
    "parameters": {
        "type": "object",
        "properties": {
            "docx": {
                "type": "string",
                "description": "Path to the .docx file.",
            },
            "table_index": {
                "type": "integer",
                "description": "Table index (0-based).",
            },
            "near_text": {
                "type": "string",
                "description": "Find the table nearest to this text snippet.",
            },
            "mapping": {
                "type": "string",
                "description": "JSON object mapping column names to values, e.g. '{\"Company\":\"Acme Ltd\"}'.",
            },
            "output": {
                "type": "string",
                "description": "Output path (default: overwrite input).",
            },
        },
        "required": ["docx", "mapping"],
    },
}


def _handle_fill_table(args: dict, **kwargs) -> str:
    from lex_docx.table_ops import fill_table
    mapping = json.loads(args["mapping"])
    kwargs_in = {"mapping": mapping}
    if "table_index" in args and args["table_index"] is not None:
        kwargs_in["table_index"] = args["table_index"]
    if "near_text" in args and args["near_text"] is not None:
        kwargs_in["near_text"] = args["near_text"]
    result = fill_table(args["docx"], **kwargs_in)
    return tool_result(result)


LEX_DOCX_TOC_SCHEMA = {
    "name": "lex_docx_toc",
    "description": (
        "Generate or refresh a Table of Contents field.  'generate' inserts "
        "a TOC field after the document title based on Heading 1-3 styles.  "
        "'refresh' updates an existing TOC (requires Word to render; this "
        "updates the field code).  Run this as the final step before issuing "
        "an execution version."
    ),
    "parameters": {
        "type": "object",
        "properties": {
            "docx": {
                "type": "string",
                "description": "Path to the .docx file.",
            },
            "action": {
                "type": "string",
                "description": "'generate' (default) or 'refresh'.",
            },
            "levels": {
                "type": "string",
                "description": "Heading level range, e.g. '1-3' (default).",
            },
            "output": {
                "type": "string",
                "description": "Output path (default: overwrite input).",
            },
        },
        "required": ["docx"],
    },
}


def _handle_toc(args: dict, **kwargs) -> str:
    from lex_docx.toc_ops import toc_generate, toc_refresh
    action = args.get("action", "generate")
    if action == "refresh":
        result = toc_refresh(args["docx"], out=args.get("output"))
    else:
        levels = args.get("levels", "1-3")
        lo, hi = levels.split("-")
        result = toc_generate(
            args["docx"],
            level_from=int(lo),
            level_to=int(hi),
            out=args.get("output"),
        )
    return tool_result(result)


# ---------------------------------------------------------------------------
# Registration
# ---------------------------------------------------------------------------

# ── Standalone registration (first tool) — required for AST-based tool discovery ──
# The discovery scanner in tools/registry.py looks for top-level
# ``registry.register(...)`` call expressions.  It does NOT recurse into loops.
# We register the first tool here so the module is picked up; the remaining 29
# are registered via the loop below.
registry.register(
    name="lex_docx_stats",
    toolset="lex-docx",
    schema=LEX_DOCX_STATS_SCHEMA,
    handler=_handle_stats,
    check_fn=_check_lex_docx,
    description=LEX_DOCX_STATS_SCHEMA.get("description", ""),
    emoji="",
)

# ── Remaining tools ─────────────────────────────────────────────────────────
_TOOLS = [
    # Inspect
    ("lex_docx_export_structure", "lex-docx", LEX_DOCX_EXPORT_STRUCTURE_SCHEMA, _handle_export_structure),
    ("lex_docx_para_query",       "lex-docx", LEX_DOCX_PARA_QUERY_SCHEMA,       _handle_para_query),
    ("lex_docx_extract_table",    "lex-docx", LEX_DOCX_EXTRACT_TABLE_SCHEMA,    _handle_extract_table),
    ("lex_docx_table_inspect",    "lex-docx", LEX_DOCX_TABLE_INSPECT_SCHEMA,    _handle_table_inspect),
    ("lex_docx_tc_list",          "lex-docx", LEX_DOCX_TC_LIST_SCHEMA,          _handle_tc_list),
    ("lex_docx_comment_list",     "lex-docx", LEX_DOCX_COMMENT_LIST_SCHEMA,     _handle_comment_list),
    ("lex_docx_footer_audit",     "lex-docx", LEX_DOCX_FOOTER_AUDIT_SCHEMA,     _handle_footer_audit),
    ("lex_docx_numbering_inspect","lex-docx", LEX_DOCX_NUMBERING_INSPECT_SCHEMA,_handle_numbering_inspect),
    # Edit
    ("lex_docx_insert",           "lex-docx", LEX_DOCX_INSERT_SCHEMA,           _handle_insert),
    ("lex_docx_replace",          "lex-docx", LEX_DOCX_REPLACE_SCHEMA,          _handle_replace),
    ("lex_docx_delete",           "lex-docx", LEX_DOCX_DELETE_SCHEMA,           _handle_delete),
    # Review
    ("lex_docx_lint",             "lex-docx", LEX_DOCX_LINT_SCHEMA,             _handle_lint),
    ("lex_docx_doctor",           "lex-docx", LEX_DOCX_DOCTOR_SCHEMA,           _handle_doctor),
    ("lex_docx_review_stats",     "lex-docx", LEX_DOCX_REVIEW_STATS_SCHEMA,     _handle_review_stats),
    # Finalize
    ("lex_docx_tc_accept",        "lex-docx", LEX_DOCX_TC_ACCEPT_SCHEMA,        _handle_tc_accept),
    ("lex_docx_tc_reject",        "lex-docx", LEX_DOCX_TC_REJECT_SCHEMA,        _handle_tc_reject),
    ("lex_docx_comment_clean",    "lex-docx", LEX_DOCX_COMMENT_CLEAN_SCHEMA,    _handle_comment_clean),
    ("lex_docx_header_clean",     "lex-docx", LEX_DOCX_HEADER_CLEAN_SCHEMA,     _handle_header_clean),
    ("lex_docx_clean",            "lex-docx", LEX_DOCX_CLEAN_SCHEMA,            _handle_clean),
    # Format
    ("lex_docx_highlight",        "lex-docx", LEX_DOCX_HIGHLIGHT_SCHEMA,        _handle_highlight),
    ("lex_docx_set_outline_level","lex-docx", LEX_DOCX_SET_OUTLINE_LEVEL_SCHEMA,_handle_set_outline_level),
    ("lex_docx_bold_terms",       "lex-docx", LEX_DOCX_BOLD_TERMS_SCHEMA,       _handle_bold_terms),
    ("lex_docx_format_brush",     "lex-docx", LEX_DOCX_FORMAT_BRUSH_SCHEMA,     _handle_format_brush),
    ("lex_docx_cleanup",          "lex-docx", LEX_DOCX_CLEANUP_SCHEMA,          _handle_cleanup),
    ("lex_docx_format_table",     "lex-docx", LEX_DOCX_FORMAT_TABLE_SCHEMA,     _handle_format_table),
    # Generate
    ("lex_docx_create",           "lex-docx", LEX_DOCX_CREATE_SCHEMA,           _handle_create),
    ("lex_docx_new_table",        "lex-docx", LEX_DOCX_NEW_TABLE_SCHEMA,        _handle_new_table),
    ("lex_docx_fill_table",       "lex-docx", LEX_DOCX_FILL_TABLE_SCHEMA,       _handle_fill_table),
    ("lex_docx_toc",              "lex-docx", LEX_DOCX_TOC_SCHEMA,              _handle_toc),
]

for _name, _toolset, _schema, _handler in _TOOLS:
    registry.register(
        name=_name,
        toolset=_toolset,
        schema=_schema,
        handler=_handler,
        check_fn=_check_lex_docx,
        description=_schema.get("description", ""),
        emoji="",
    )


def reload_lex_docx_tools() -> dict:
    """Hot-reload all lex_docx tools by re-importing this module.

    Deregisters existing lex-docx tools, forces a fresh import of the
    module (which re-runs all ``registry.register()`` calls), and clears
    the check_fn TTL cache so availability probes see the new code.

    Returns a dict with ``deregistered``, ``reregistered`` counts.
    """
    before = set(registry.get_tool_names_for_toolset("lex-docx"))

    # Drop every existing lex-docx tool entry
    for name in list(before):
        registry.deregister(name)

    # Re-import this module — runs all module-level registry.register() again
    import importlib
    import tools.lex_docx_tool
    importlib.reload(tools.lex_docx_tool)

    # Drop the check_fn cache so the module's new check_fn takes effect
    invalidate_check_fn_cache()

    after = set(registry.get_tool_names_for_toolset("lex-docx"))
    return {
        "deregistered": len(before),
        "reregistered": len(after),
        "tools": sorted(after),
    }
